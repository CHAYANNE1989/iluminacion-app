"""
generar_word.py  —  LuxOMeter PRO / RETILAP 2024
Genera el informe Word usando la plantilla de la ARL seleccionada.
Reemplaza campos amarillos, tabla de equipo, Tabla 2, Tabla 4 y Gráfica 1.
"""
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Rutas de plantillas ───────────────────────────────────────────────────────
PLANTILLA_PATH = "INFORME_PREFORMA.docx"

# ── Colores ───────────────────────────────────────────────────────────────────
AZ_OSC = RGBColor(0x1A, 0x3A, 0x5C)
AZ_CLA = RGBColor(0xD6, 0xE4, 0xF0)
VERDE  = RGBColor(0x27, 0xAE, 0x60)
ROJO   = RGBColor(0xE7, 0x4C, 0x3C)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO  = RGBColor(0x00, 0x00, 0x00)
GRIS   = RGBColor(0xF0, 0xF4, 0xF8)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hx = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hx)
    tcPr.append(shd)


def _txt(cell, text, bold=False, sz=8, color=NEGRO,
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.clear(); p.alignment = align
    r = p.add_run(str(text) if text else "")
    r.bold = bold; r.italic = italic
    r.font.size = Pt(sz); r.font.color.rgb = color


def _borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{s}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0');   e.set(qn('w:color'), 'AACCEE')
        b.append(e)
    tblPr.append(b)


def _is_yellow(run):
    try:    return run.font.highlight_color is not None
    except: return False


def _img_buf(pil_img):
    buf = io.BytesIO(); pil_img.save(buf, format="PNG"); buf.seek(0); return buf


def _mes_texto(fecha_str):
    meses = {'01':'ENERO','02':'FEBRERO','03':'MARZO','04':'ABRIL',
             '05':'MAYO','06':'JUNIO','07':'JULIO','08':'AGOSTO',
             '09':'SEPTIEMBRE','10':'OCTUBRE','11':'NOVIEMBRE','12':'DICIEMBRE'}
    try:
        p = fecha_str.split('/')
        return f"{meses.get(p[1], p[1])} - {p[2]}"
    except:
        return fecha_str


def _reemplazar_texto_amarillo(para, nuevo):
    """Reemplaza todos los runs amarillos de un párrafo."""
    primero = True
    for run in para.runs:
        if _is_yellow(run):
            if primero and run.text.strip() not in ('', ',', ', '):
                run.text = nuevo; primero = False
            elif not primero:
                run.text = ""
            run.font.highlight_color = None


def _reemplazar_en_parrafo(para, viejo, nuevo):
    """Reemplaza texto literal en todos los runs de un párrafo."""
    for run in para.runs:
        if viejo in run.text:
            run.text = run.text.replace(viejo, nuevo)


# ── Grafica de barras ─────────────────────────────────────────────────────────

def _generar_grafica_bytes(mediciones):
    try:
        import matplotlib; matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        total      = len(mediciones)
        conformes  = sum(1 for m in mediciones if "✅" in str(m.get("resultado","")))
        deficientes = total - conformes
        if total == 0: return None

        pct_conf = round(conformes/total*100, 1)
        pct_def  = round(deficientes/total*100, 1)

        fig, ax = plt.subplots(figsize=(7, 3), facecolor='#f8fafc')
        ax.set_facecolor('#f8fafc')
        bars = ax.barh([1,0], [pct_conf, pct_def],
                       color=['#27ae60','#e74c3c'], height=0.5,
                       edgecolor='white', linewidth=1.5)
        for bar, val, n in zip(bars, [pct_conf,pct_def], [conformes,deficientes]):
            ax.text(val/2, bar.get_y()+bar.get_height()/2,
                    f"{val}%  ({n} pts)",
                    ha='center', va='center', fontsize=11, fontweight='bold', color='white')
        ax.set_yticks([1,0])
        ax.set_yticklabels(['Adecuados','Deficientes'], fontsize=11,
                           fontweight='bold', color='#1a3a5c')
        ax.set_xlim(0, 115)
        ax.set_xlabel('Porcentaje (%)', fontsize=9, color='#475569')
        ax.set_title(f'Conformidad Lumínica — {total} puntos evaluados',
                     fontsize=11, fontweight='bold', color='#1a3a5c', pad=10)
        ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
        ax.spines['left'].set_visible(False)
        ax.tick_params(axis='x', colors='#94a3b8'); ax.tick_params(axis='y', left=False)
        ax.xaxis.grid(True, linestyle='--', alpha=0.4, color='#cbd5e1'); ax.set_axisbelow(True)
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='PNG', bbox_inches='tight', dpi=140, facecolor='#f8fafc')
        plt.close(fig); buf.seek(0); return buf.getvalue()
    except: return None


# ── Tabla de resultados ───────────────────────────────────────────────────────

def _insertar_tabla_resultados(doc, para_ref, mediciones):
    HEADS = [
        "N°\nMed","Puesto de trabajo\no Área evaluada","Descripción",
        "E\nMIN\n(lx)","E\nMAX\n(lx)","Promedio\nmedido\n(lx)",
        "Valor\nUo","Interp.\nUo","Tipo de Área\nRETILAP",
        "Em\nrec.\n(lx)","Interpretación\ndel Nivel de\nIluminancia",
        "Observaciones /\nRecomendaciones",
    ]
    CW  = [0.9,3.2,2.8,1.1,1.1,1.3,1.1,1.1,3.2,1.1,2.0,3.5]
    NC  = len(HEADS)

    tbl = doc.add_table(rows=1, cols=NC)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER; _borders(tbl)
    hr = tbl.rows[0]
    for ci,(h,cw) in enumerate(zip(HEADS,CW)):
        cell=hr.cells[ci]; cell.width=Cm(cw)
        _bg(cell,AZ_OSC); _txt(cell,h,bold=True,color=BLANCO,sz=7)

    for idx_m,m in enumerate(mediciones):
        conf_m = "✅" in str(m.get("resultado",""))
        rbg    = GRIS if idx_m%2==0 else RGBColor(0xFF,0xFF,0xFF)
        m1=m.get("med1",0) or 0; m2=m.get("med2",0) or 0
        m3=m.get("med3",0) or 0; m4=m.get("med4",0) or 0
        vals=[v for v in[m1,m2,m3,m4] if v>0]
        e_min = m.get("e_min") or (round(min(vals),1) if vals else "")
        e_max = m.get("e_max") or (round(max(vals),1) if vals else "")
        desc  = (f"Tipo Ilum.: {m.get('tipo_iluminacion','')}\n"
                 f"Lámpara: {m.get('tipo_lampara','')}\n"
                 f"Ubic.: {m.get('ubicacion_luminaria','')}\n"
                 f"Ctrl. Luz Nat.: {m.get('control_luz_natural','')}\n"
                 f"Altura (m): {m.get('altura_luminaria','')}")
        obs   = (f"Obs.: {m.get('nota','')}\n"
                 f"Rec.: {m.get('recomendacion','')}")
        vr = [str(m.get("num","")),
              str(m.get("puesto_evaluado","") or m.get("area","")),
              desc, str(e_min), str(e_max),
              str(m.get("promedio","")), str(m.get("uo_calc","")),
              str(m.get("interpretacion_uo","")), str(m.get("area","")),
              str(m.get("em_req","")),
              "ADECUADO" if conf_m else "DEFICIENTE", obs]
        dr = tbl.add_row()
        for ci,(val,cw) in enumerate(zip(vr,CW)):
            cell=dr.cells[ci]; cell.width=Cm(cw)
            if ci==NC-2:
                _bg(cell,VERDE if conf_m else ROJO)
                _txt(cell,val,bold=True,color=BLANCO,sz=7)
            else:
                _bg(cell,rbg)
                al=WD_ALIGN_PARAGRAPH.LEFT if ci in(1,2,8,NC-1) else WD_ALIGN_PARAGRAPH.CENTER
                _txt(cell,val,sz=7,align=al)

    # Mover tabla justo después del párrafo de referencia
    para_ref._p.addnext(tbl._tbl)


# ── Tabla 2 RETILAP con áreas reales ─────────────────────────────────────────

def _actualizar_tabla2_retilap(tabla, mediciones):
    areas_usadas = []
    vistas = set()
    for m in mediciones:
        area = m.get("area",""); em = m.get("em_req",""); uo = m.get("uo_min","")
        if area and area not in vistas:
            vistas.add(area); areas_usadas.append((area, str(em), str(uo)))
    if not areas_usadas: return

    tbl = tabla._tbl; filas = tbl.findall(qn('w:tr'))
    for fila in filas[1:]: tbl.remove(fila)

    for area, em, uo in areas_usadas:
        tr = OxmlElement('w:tr')
        for ci, val in enumerate([area, f"{em} lx" if em else "", uo]):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            try:
                orig_w = filas[0].findall(qn('w:tc'))[ci].find(qn('w:tcPr')).find(qn('w:tcW'))
                if orig_w is not None:
                    nw = OxmlElement('w:tcW')
                    nw.set(qn('w:w'), orig_w.get(qn('w:w')))
                    nw.set(qn('w:type'), orig_w.get(qn('w:type')))
                    tcPr.append(nw)
            except: pass
            tc.append(tcPr)
            p = OxmlElement('w:p'); pPr = OxmlElement('w:pPr')
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center' if ci>0 else 'left')
            pPr.append(jc); p.append(pPr)
            r = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')
            rPr.append(sz); r.append(rPr)
            t_e = OxmlElement('w:t'); t_e.text = val; r.append(t_e); p.append(r); tc.append(p); tr.append(tc)
        tbl.append(tr)


# ── Actualizador de tabla de equipo (T0) ──────────────────────────────────────

def _actualizar_equipo(tabla, equipo):
    """Rellena marca, modelo y serie en la tabla de instrumento."""
    if len(tabla.rows) < 2: return
    fila = tabla.rows[1]
    campos = ['instrumento', 'marca', 'modelo', 'serie']
    for ci, campo in enumerate(campos):
        if ci < len(fila.cells) and campo in equipo:
            cell = fila.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(equipo[campo]))
            run.font.size = Pt(9)


# ── Lógica de reemplazo por ARL ───────────────────────────────────────────────

def _rellenar_plantilla(doc, arl, general, mediciones, equipo):
    """Reemplaza todos los campos amarillos según la ARL."""
    empresa   = general.get("nombre_empresa","").upper()
    ciudad    = general.get("sede","").upper()
    fecha     = general.get("fecha", datetime.now().strftime('%d/%m/%Y'))
    mes_anio  = _mes_texto(fecha).upper()
    nit       = general.get("nit","")
    direccion = general.get("direccion","")
    telefono  = general.get("telefono","")
    resp_emp  = general.get("responsable_empresa","")
    resp_hig  = general.get("responsable_higienista","")
    orden     = general.get("numero_orden","")
    paras     = doc.paragraphs

    # Textos placeholder comunes en todas las plantillas
    EMPRESA_PLACEHOLDER = ["INDEPENDIENTE SANTA FE", "INDEPENDIENTE SANTAFE"]

    # ── Reemplazar nombre empresa y ciudad en TODO el documento ───────────────
    for para in paras:
        for ph in EMPRESA_PLACEHOLDER:
            if ph in para.text:
                _reemplazar_en_parrafo(para, ph, empresa)
        if 'BOGOTA' in para.text and not any(_is_yellow(r) for r in para.runs):
            _reemplazar_en_parrafo(para, 'BOGOTA', ciudad)
        if 'MARZO - 2026' in para.text and not any(_is_yellow(r) for r in para.runs):
            _reemplazar_en_parrafo(para, 'MARZO - 2026', mes_anio)

    if arl == "Positiva":
        # P6=empresa, P7=ciudad, P30=ciudad+mes, P65=empresa en objetivo
        _map = {6: empresa, 7: ciudad}
        for idx, val in _map.items():
            if idx < len(paras):
                for run in paras[idx].runs:
                    if _is_yellow(run): run.text = val if run.text.strip() else ""; run.font.highlight_color = None

        # P30: ciudad + mes/año
        if 30 < len(paras):
            ciudad_done = False
            for run in paras[30].runs:
                if _is_yellow(run):
                    if not ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = ciudad; ciudad_done = True
                    elif ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = mes_anio
                    run.font.highlight_color = None

    elif arl == "Colmena":
        # P10=empresa, P11=ciudad, P29=ciudad+mes
        # P34=NIT, P36=empresa, P38=dirección, P40=tel, P42=resp_emp, P44=resp_hig
        # P46=ciudad (centro trabajo), P48=ciudad (ejecución)
        _map_yellow = {
            10: empresa, 11: ciudad,
            36: empresa, 38: direccion, 40: telefono,
            42: resp_emp, 44: resp_hig,
            46: ciudad, 48: ciudad,
        }
        for idx, val in _map_yellow.items():
            if idx < len(paras):
                for run in paras[idx].runs:
                    if _is_yellow(run):
                        run.text = val if run.text.strip() not in ('',':',': ') else run.text
                        run.font.highlight_color = None
        # P34: NIT (formato "NIT EMPRESA: XXXX")
        if 34 < len(paras):
            for run in paras[34].runs:
                if _is_yellow(run) and run.text.strip():
                    run.text = nit; run.font.highlight_color = None
        # P29: ciudad + mes
        if 29 < len(paras):
            ciudad_done = False
            for run in paras[29].runs:
                if _is_yellow(run):
                    if not ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = ciudad; ciudad_done = True
                    elif ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = mes_anio
                    run.font.highlight_color = None

    elif arl == "Bolívar":
        # P21=CRONOGRAMA(orden), P22=SECUENCIA(orden)
        # P24=empresa, P25=ciudad, P26=NIT, P29=ciudad+mes
        if 21 < len(paras):
            for run in paras[21].runs:
                if _is_yellow(run) and run.text.strip(): run.text = orden; run.font.highlight_color = None
        if 22 < len(paras):
            for run in paras[22].runs:
                if _is_yellow(run) and run.text.strip(): run.text = orden; run.font.highlight_color = None
        _map_yellow = {24: empresa, 25: ciudad}
        for idx, val in _map_yellow.items():
            if idx < len(paras):
                for run in paras[idx].runs:
                    if _is_yellow(run): run.text = val if run.text.strip() else ""; run.font.highlight_color = None
        # P26: NIT (formato "NIT XXXX")
        if 26 < len(paras):
            for run in paras[26].runs:
                if _is_yellow(run) and run.text.strip(): run.text = nit; run.font.highlight_color = None
        # P29: ciudad + mes
        if 29 < len(paras):
            ciudad_done = False
            for run in paras[29].runs:
                if _is_yellow(run):
                    if not ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = ciudad; ciudad_done = True
                    elif ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = mes_anio
                    run.font.highlight_color = None

    elif arl == "Sura":
        # P9=empresa, P10=ciudad, P28=ciudad+mes
        _map_yellow = {9: empresa, 10: ciudad}
        for idx, val in _map_yellow.items():
            if idx < len(paras):
                for run in paras[idx].runs:
                    if _is_yellow(run): run.text = val if run.text.strip() else ""; run.font.highlight_color = None
        if 28 < len(paras):
            ciudad_done = False
            for run in paras[28].runs:
                if _is_yellow(run):
                    if not ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = ciudad; ciudad_done = True
                    elif ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = mes_anio
                    run.font.highlight_color = None

    elif arl == "AXA Colpatria":
        # P7=empresa, P8=ciudad, P25=ciudad+mes
        # P55=orden, P64=ciudad, P67=empresa (en "Trabajadores de..."), P70=mes, P73=mes
        _map_yellow = {7: empresa, 8: ciudad, 64: ciudad}
        for idx, val in _map_yellow.items():
            if idx < len(paras):
                for run in paras[idx].runs:
                    if _is_yellow(run): run.text = val if run.text.strip() else ""; run.font.highlight_color = None
        if 55 < len(paras):
            for run in paras[55].runs:
                if _is_yellow(run) and run.text.strip(): run.text = orden; run.font.highlight_color = None
        if 67 < len(paras):
            for run in paras[67].runs:
                if _is_yellow(run): run.text = f"Trabajadores de la empresa {empresa}"; run.font.highlight_color = None
        for idx in [70, 73]:
            if idx < len(paras):
                for run in paras[idx].runs:
                    if _is_yellow(run) and run.text.strip(): run.text = mes_anio; run.font.highlight_color = None
        # P25: ciudad + mes
        if 25 < len(paras):
            ciudad_done = False
            for run in paras[25].runs:
                if _is_yellow(run):
                    if not ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = ciudad; ciudad_done = True
                    elif ciudad_done and run.text.strip() not in ('',',',', '):
                        run.text = mes_anio
                    run.font.highlight_color = None


# ── FUNCIÓN PRINCIPAL ─────────────────────────────────────────────────────────

def generar_informe_word(general: dict, mediciones: list,
                         plano_imgs: dict = None,
                         arl: str = "Positiva",
                         plantillas_arl: dict = None) -> bytes:
    """
    Genera informe Word usando la plantilla de la ARL seleccionada.
    """
    if plantillas_arl is None:
        plantillas_arl = {
            "Positiva":      "INFORME_PREFORMA.docx",
            "Colmena":       "INFORME_COLMENA.docx",
            "Bolívar":       "INFORME_BOLIVAR.docx",
            "AXA Colpatria": "INFORME_AXA.docx",
            "Sura":          "INFORME_SURA.docx",
        }

    plantilla = plantillas_arl.get(arl, PLANTILLA_PATH)
    equipo    = general.get("equipo", {})

    # ── Cargar plantilla ────────────────────────────────────────────────────
    if os.path.exists(plantilla):
        doc = Document(plantilla)
    elif os.path.exists(PLANTILLA_PATH):
        doc = Document(PLANTILLA_PATH)
    else:
        return _generar_sin_plantilla(general, mediciones, plano_imgs, arl)

    # ── 1. Rellenar campos amarillos según ARL ──────────────────────────────
    _rellenar_plantilla(doc, arl, general, mediciones, equipo)

    # ── 2. Actualizar tabla de equipo (T0) ──────────────────────────────────
    if doc.tables and equipo:
        _actualizar_equipo(doc.tables[0], equipo)

    # ── 3. Actualizar Tabla 2 RETILAP con áreas reales ──────────────────────
    if len(doc.tables) > 1 and mediciones:
        _actualizar_tabla2_retilap(doc.tables[1], mediciones)

    # ── 4. Insertar tabla de resultados (busca "Tabla 4" o "Tabla 5") ───────
    paras = doc.paragraphs
    tabla_ref = None
    for para in paras:
        txt = para.text.strip()
        if txt.startswith('Tabla 4') or txt.startswith('Tabla 5'):
            tabla_ref = para; break
    if tabla_ref and mediciones:
        _insertar_tabla_resultados(doc, tabla_ref, mediciones)

    # ── 5. Insertar gráfica en "Grafica 1" ──────────────────────────────────
    for para in paras:
        if para.text.strip().startswith('Grafica 1') or para.text.strip().startswith('Gráfica 1'):
            graf_bytes = _generar_grafica_bytes(mediciones)
            if graf_bytes:
                try:
                    p_graf = doc.add_paragraph()
                    p_graf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_graf.add_run().add_picture(io.BytesIO(graf_bytes), width=Cm(14))
                    para._p.addnext(p_graf._p)
                except: pass
            break

    # ── 6. Planos ────────────────────────────────────────────────────────────
    if plano_imgs:
        for pln_nombre, pln_img in plano_imgs.items():
            if pln_img:
                try:
                    p = doc.add_paragraph()
                    r = p.add_run(f"Plano: {pln_nombre}")
                    r.bold = True; r.font.color.rgb = AZ_OSC
                    doc.add_picture(_img_buf(pln_img), width=Cm(22))
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f"(Error plano: {e})")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()


# ── FALLBACK sin plantilla ────────────────────────────────────────────────────

def _generar_sin_plantilla(general, mediciones, plano_imgs, arl=""):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(29.7); sec.page_height=Cm(21.0)
    sec.left_margin=sec.right_margin=Cm(2.5)
    sec.top_margin=sec.bottom_margin=Cm(2.0)

    empresa   = general.get("nombre_empresa","").upper()
    nit       = general.get("nit","")
    ciudad    = general.get("sede","")
    fecha     = general.get("fecha","")
    orden     = general.get("numero_orden","")
    higienista= general.get("responsable_higienista","")
    resolucion= general.get("resolucion","")
    direccion = general.get("direccion","")
    resp_emp  = general.get("responsable_empresa","")
    mes_anio  = _mes_texto(fecha).upper()
    equipo    = general.get("equipo", {})

    def _h1(txt):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(txt); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=AZ_OSC

    def _h2(txt):
        p=doc.add_paragraph(); r=p.add_run(txt)
        r.bold=True; r.font.size=Pt(12); r.font.color.rgb=AZ_OSC
        p.paragraph_format.space_before=Pt(12)

    def _sep():
        p=doc.add_paragraph("─"*80); p.runs[0].font.color.rgb=AZ_CLA; p.runs[0].font.size=Pt(8)

    def _norm(txt, bold=False, sz=10, center=False):
        p=doc.add_paragraph()
        if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(txt); r.bold=bold; r.font.size=Pt(sz)

    # Portada
    doc.add_paragraph()
    _h1("EVALUACIÓN DE INTENSIDAD LUMÍNICA")
    doc.add_paragraph()
    _norm(empresa, bold=True, sz=22, center=True)
    _norm(f"{ciudad}  ·  {mes_anio}", sz=12, center=True)
    if arl: _norm(f"ARL: {arl}", sz=11, center=True)
    doc.add_paragraph(); _sep()
    _norm(f"Orden: {orden}  |  Fecha: {fecha}", sz=9, center=True)
    _sep(); doc.add_paragraph()
    _norm("HISIG CONSULTORÍA S.A.S", bold=True, sz=11, center=True)
    _norm("ESPECIALISTAS EN HIGIENE OCUPACIONAL", sz=9, center=True)
    _norm(f"Resolución No. {resolucion}", sz=9, center=True)
    doc.add_page_break()

    # Datos empresa
    _h2("1. DATOS DE LA EMPRESA"); _sep()
    ficha=doc.add_table(rows=5,cols=4); ficha.alignment=WD_TABLE_ALIGNMENT.LEFT; _borders(ficha)
    cw_f=[3.5,5.0,3.5,5.0]
    fd=[["Empresa:",empresa,"NIT:",nit],
        ["Dirección:",direccion,"Ciudad:",ciudad],
        ["Teléfono:",general.get("telefono",""),"Responsable:",resp_emp],
        ["Higienista:",higienista,"Lic. SST:",resolucion],
        ["N° Orden:",orden,"ARL:",arl]]
    for ri,fila in enumerate(fd):
        bg=GRIS if ri%2==0 else RGBColor(0xFF,0xFF,0xFF)
        for ci,val in enumerate(fila):
            cell=ficha.cell(ri,ci); cell.width=Cm(cw_f[ci]); _bg(cell,bg)
            _txt(cell,val,bold=(ci%2==0),sz=9,
                 color=AZ_OSC if ci%2==0 else NEGRO,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_paragraph()

    # Equipo
    _h2("2. EQUIPO DE MEDICIÓN"); _sep()
    eq_tbl=doc.add_table(rows=2,cols=4); eq_tbl.alignment=WD_TABLE_ALIGNMENT.LEFT; _borders(eq_tbl)
    for ci,h in enumerate(["Instrumento","Marca","Modelo","Serie"]):
        cell=eq_tbl.cell(0,ci); _bg(cell,AZ_OSC); _txt(cell,h,bold=True,color=BLANCO,sz=9)
    vals_eq=["Luxómetro",
             equipo.get("marca",""), equipo.get("modelo",""), equipo.get("serie","")]
    for ci,v in enumerate(vals_eq):
        cell=eq_tbl.cell(1,ci); _bg(cell,GRIS); _txt(cell,v,sz=9)
    doc.add_paragraph()

    # Resumen ejecutivo
    _h2("3. RESUMEN EJECUTIVO"); _sep()
    tot=len(mediciones); conf=sum(1 for m in mediciones if "✅" in str(m.get("resultado",""))); defic=tot-conf
    pct=round(conf/tot*100,1) if tot>0 else 0
    rs=doc.add_table(rows=2,cols=4); rs.alignment=WD_TABLE_ALIGNMENT.CENTER; _borders(rs)
    for ci,h in enumerate(["Total puntos","Adecuados","Deficientes","% Adecuados"]):
        cell=rs.cell(0,ci); _bg(cell,AZ_OSC); _txt(cell,h,bold=True,color=BLANCO,sz=10)
    for ci,v in enumerate([str(tot),str(conf),str(defic),f"{pct}%"]):
        cell=rs.cell(1,ci); _bg(cell,GRIS)
        col=(VERDE if pct>=80 else ROJO) if ci==3 else NEGRO
        _txt(cell,v,bold=(ci==3),color=col,sz=11)
    doc.add_paragraph()
    graf_bytes=_generar_grafica_bytes(mediciones)
    if graf_bytes:
        p_g=doc.add_paragraph(); p_g.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_g.add_run().add_picture(io.BytesIO(graf_bytes),width=Cm(14))
    doc.add_page_break()

    # Planos
    if plano_imgs:
        _h2("4. PLANOS DE MEDICIÓN"); _sep()
        for pln,img in plano_imgs.items():
            _norm(f"Plano: {pln}",bold=True,sz=10)
            if img:
                try: doc.add_picture(_img_buf(img),width=Cm(22))
                except: pass
            doc.add_paragraph()
        doc.add_page_break()

    # Tabla de resultados
    _h2("5. RESULTADOS — RETILAP 2024"); _sep()
    HEADS=["N°\nMed","Puesto de trabajo\no Área evaluada","Descripción",
           "E\nMIN\n(lx)","E\nMAX\n(lx)","Promedio\nmedido\n(lx)",
           "Valor\nUo","Interp.\nUo","Tipo de Área\nRETILAP",
           "Em\nrec.\n(lx)","Interpretación\ndel Nivel de\nIluminancia",
           "Observaciones /\nRecomendaciones"]
    CW=[0.9,3.2,2.8,1.1,1.1,1.3,1.1,1.1,3.2,1.1,2.0,3.5]; NC=len(HEADS)
    tbl=doc.add_table(rows=1,cols=NC); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; _borders(tbl)
    for ci,(h,cw) in enumerate(zip(HEADS,CW)):
        cell=tbl.rows[0].cells[ci]; cell.width=Cm(cw); _bg(cell,AZ_OSC); _txt(cell,h,bold=True,color=BLANCO,sz=7)
    for idx_m,m in enumerate(mediciones):
        conf_m="✅" in str(m.get("resultado",""))
        rbg=GRIS if idx_m%2==0 else RGBColor(0xFF,0xFF,0xFF)
        m1=m.get("med1",0) or 0; m2=m.get("med2",0) or 0
        m3=m.get("med3",0) or 0; m4=m.get("med4",0) or 0
        vals=[v for v in[m1,m2,m3,m4] if v>0]
        e_min=m.get("e_min") or (round(min(vals),1) if vals else "")
        e_max=m.get("e_max") or (round(max(vals),1) if vals else "")
        desc=(f"Tipo Ilum.: {m.get('tipo_iluminacion','')}\nLámpara: {m.get('tipo_lampara','')}\n"
              f"Ubic.: {m.get('ubicacion_luminaria','')}\nCtrl. Luz Nat.: {m.get('control_luz_natural','')}\n"
              f"Altura (m): {m.get('altura_luminaria','')}")
        obs=f"Obs.: {m.get('nota','')}\nRec.: {m.get('recomendacion','')}"
        vr=[str(m.get("num","")),str(m.get("puesto_evaluado","") or m.get("area","")),
            desc,str(e_min),str(e_max),str(m.get("promedio","")),
            str(m.get("uo_calc","")),str(m.get("interpretacion_uo","")),
            str(m.get("area","")),str(m.get("em_req","")),
            "ADECUADO" if conf_m else "DEFICIENTE",obs]
        dr=tbl.add_row()
        for ci,(val,cw) in enumerate(zip(vr,CW)):
            cell=dr.cells[ci]; cell.width=Cm(cw)
            if ci==NC-2: _bg(cell,VERDE if conf_m else ROJO); _txt(cell,val,bold=True,color=BLANCO,sz=7)
            else:
                _bg(cell,rbg)
                al=WD_ALIGN_PARAGRAPH.LEFT if ci in(1,2,8,NC-1) else WD_ALIGN_PARAGRAPH.CENTER
                _txt(cell,val,sz=7,align=al)

    doc.add_paragraph(); _sep()
    p_pie=doc.add_paragraph(); p_pie.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r_pie=p_pie.add_run(f"RETILAP 2024  ·  {empresa}  ·  {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    r_pie.font.size=Pt(7); r_pie.font.color.rgb=RGBColor(0x80,0x80,0x80)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()
